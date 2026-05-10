"""DOCX and PDF output with template support.

Both formats support {{field_name}} substitution markers. Templates can be:
  - The bundled default (whisperapp/templates/default.html for PDF,
    auto-generated DOCX for Word)
  - A user-supplied file set via config.output_template_docx /
    config.output_template_pdf
  - A file dropped into ~/.whisperapp/templates/

Available markers
-----------------
{{filename}}            Source file name
{{filepath}}            Full source path
{{date}}                Transcription date (YYYY-MM-DD)
{{datetime}}            Date and time (YYYY-MM-DD HH:MM)
{{duration}}            Audio duration (HH:MM:SS or MM:SS)
{{model}}               Whisper model used
{{language}}            Detected language
{{transcript}}          Plain transcript (no speaker labels)
{{diarized_transcript}} Transcript with [Speaker]: prefixes
{{speakers}}            Comma-separated speaker names
{{word_count}}          Total word count
{{meeting_notes}}       AI-generated meeting notes (blank if not run)
{{segments}}            Timestamped segments, one per line
{{source_hash}}         SHA-256 hex digest of the source file
{{file_size_bytes}}     Source file size in bytes
{{legal_transcript}}    Line-numbered transcript (L0001 HH:MM:SS [Speaker] text)
"""

from __future__ import annotations
import datetime
import logging
import re
from pathlib import Path
from typing import Optional

log = logging.getLogger(__name__)

_PKG_TEMPLATES = Path(__file__).parent / "templates"


# ── Context ────────────────────────────────────────────────────────────────

def build_context(result: dict, source_file: str) -> dict:
    meta = result.get("source_metadata") or {}
    segments = result.get("segments", [])

    plain_lines, diarized_lines, seg_lines, legal_lines = [], [], [], []
    speakers_seen: set[str] = set()
    line_num = 0

    for seg in segments:
        line_num += 1
        text = seg.get("text", "").strip()
        speaker = seg.get("speaker", "")
        num = f"L{line_num:04d}"

        if seg.get("is_pause_marker"):
            plain_lines.append(text)
            diarized_lines.append(text)
            seg_lines.append(f"           {text}")
            legal_lines.append(f"{num}  {'':20}  {'':14}  {text}")
            continue

        plain_lines.append(text)
        diarized_lines.append(f"[{speaker}]: {text}" if speaker else text)
        if speaker:
            speakers_seen.add(speaker)

        start = _fmt_ts(seg.get("start", 0))
        end   = _fmt_ts(seg.get("end", 0))
        prefix = f"[{speaker}] " if speaker else ""
        seg_lines.append(f"{start} → {end}  {prefix}{text}")

        ts_long = _fmt_ts_long(seg.get("start", 0))
        speaker_col = f"[{speaker}]" if speaker else ""
        legal_lines.append(f"{num}  {ts_long:<10}  {speaker_col:<14}  {text}")

    word_count = sum(
        len(seg.get("text", "").split())
        for seg in segments
        if not seg.get("is_pause_marker")
    )

    return {
        "filename":            meta.get("source_file", Path(source_file).name),
        "filepath":            str(source_file),
        "date":                datetime.date.today().isoformat(),
        "datetime":            datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
        "duration":            meta.get("duration", "—"),
        "model":               result.get("model", "—"),
        "language":            result.get("language", "—"),
        "transcript":          "\n".join(plain_lines),
        "diarized_transcript": "\n".join(diarized_lines),
        "speakers":            ", ".join(sorted(speakers_seen)) if speakers_seen else "—",
        "word_count":          str(word_count),
        "meeting_notes":       result.get("meeting_notes", ""),
        "segments":            "\n".join(seg_lines),
        "source_hash":         meta.get("source_hash", "—"),
        "file_size_bytes":     meta.get("file_size_bytes", "—"),
        "legal_transcript":    "\n".join(legal_lines),
    }


def _fmt_ts(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    return f"{h:02d}:{m:02d}:{s:02d}" if h else f"{m:02d}:{s:02d}"


def _fmt_ts_long(seconds: float) -> str:
    """Always HH:MM:SS — used for legal transcripts where column alignment matters."""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def substitute(text: str, ctx: dict) -> str:
    """Replace {{field}} markers in text with context values."""
    def _replace(m):
        return ctx.get(m.group(1).strip(), m.group(0))
    return re.sub(r"\{\{(\w+)\}\}", _replace, text)


# ── DOCX ───────────────────────────────────────────────────────────────────

def _resolve_docx_template(config) -> Optional[Path]:
    for candidate in [
        config and getattr(config, "output_template_docx", ""),
        str(Path.home() / ".whisperapp" / "templates" / "custom.docx"),
    ]:
        if candidate:
            p = Path(candidate)
            if p.exists():
                return p
    return None


def _default_docx(ctx: dict):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin    = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin   = Pt(90)
        section.right_margin  = Pt(90)

    # Title
    h = doc.add_heading(ctx["filename"], 0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata table
    meta_rows = [
        ("Date",      ctx["date"]),
        ("Duration",  ctx["duration"]),
        ("Speakers",  ctx["speakers"]),
        ("Words",     ctx["word_count"]),
        ("Model",     ctx["model"]),
        ("Language",  ctx["language"]),
    ]
    tbl = doc.add_table(rows=len(meta_rows), cols=2)
    tbl.style = "Table Grid"
    for i, (label, value) in enumerate(meta_rows):
        row = tbl.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph()

    # Meeting notes
    if ctx.get("meeting_notes"):
        doc.add_heading("Meeting Notes", 1)
        doc.add_paragraph(ctx["meeting_notes"])
        doc.add_paragraph()

    # Transcript
    doc.add_heading("Transcript", 1)
    for line in ctx["diarized_transcript"].split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        # Bold the [Speaker]: prefix if present
        if line.startswith("[") and "]: " in line:
            end_bracket = line.index("]: ") + 1
            run_label = p.add_run(line[:end_bracket + 1])
            run_label.bold = True
            p.add_run(" " + line[end_bracket + 2:])
        else:
            run = p.add_run(line)
            if line.startswith("[") and line.endswith("]"):
                run.italic = True
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Footer paragraph
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.paragraph_format.space_before = Pt(12)
    run = footer_p.add_run(f"Generated by WhisperApp  ·  {ctx['filepath']}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    return doc


def _template_docx(template_path: Path, ctx: dict):
    """Load a .docx template and substitute {{field}} markers in all text."""
    from docx import Document

    doc = Document(str(template_path))

    def _sub_para(para):
        full = "".join(r.text for r in para.runs)
        if "{{" not in full:
            return
        new_text = substitute(full, ctx)
        if para.runs:
            para.runs[0].text = new_text
            for r in para.runs[1:]:
                r.text = ""

    for para in doc.paragraphs:
        _sub_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _sub_para(para)

    return doc


def write_docx(result: dict, source_file: str, out_path: Path, config=None):
    try:
        from docx import Document  # noqa: F401
    except ImportError:
        log.warning("python-docx not installed — skipping DOCX. "
                    "Install with: pip install python-docx")
        return

    ctx = build_context(result, source_file)
    template = _resolve_docx_template(config)

    doc = _template_docx(template, ctx) if template else _default_docx(ctx)
    doc.save(str(out_path))
    log.info("Wrote DOCX: %s", out_path)


def generate_docx_template(dest: Path):
    """Write an editable DOCX template to dest for the user to customise."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        log.warning("python-docx not installed — cannot generate template")
        return False

    doc = Document()
    doc.add_heading("{{filename}}", 0)

    intro = doc.add_paragraph(
        "Edit this template to match your organisation's style. "
        "WhisperApp replaces {{field_name}} markers with real values before saving."
    )
    intro.runs[0].italic = True
    intro.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    fields = [
        ("{{date}}",                "Transcription date"),
        ("{{datetime}}",            "Date and time"),
        ("{{duration}}",            "Audio duration"),
        ("{{speakers}}",            "Identified speakers (comma-separated)"),
        ("{{word_count}}",          "Total word count"),
        ("{{model}}",               "Whisper model used"),
        ("{{language}}",            "Detected language"),
        ("{{filename}}",            "Source file name"),
        ("{{filepath}}",            "Full source file path"),
        ("{{transcript}}",          "Plain transcript text"),
        ("{{diarized_transcript}}", "Transcript with [Speaker]: labels"),
        ("{{segments}}",            "Timestamped segments, one per line"),
        ("{{meeting_notes}}",       "AI meeting notes (blank if not run)"),
    ]

    doc.add_heading("Available Fields", 1)
    tbl = doc.add_table(rows=len(fields), cols=2)
    tbl.style = "Table Grid"
    for i, (marker, desc) in enumerate(fields):
        row = tbl.rows[i]
        row.cells[0].text = marker
        row.cells[1].text = desc
        for run in row.cells[0].paragraphs[0].runs:
            run.font.name = "Courier New"
            run.bold = True
    doc.add_paragraph()

    doc.add_heading("Example — Meeting Report", 1)
    doc.add_paragraph("Date: {{date}}")
    doc.add_paragraph("File: {{filename}}  ·  Duration: {{duration}}")
    doc.add_paragraph("Attendees: {{speakers}}")
    doc.add_paragraph()

    doc.add_heading("Notes", 2)
    doc.add_paragraph("{{meeting_notes}}")
    doc.add_paragraph()

    doc.add_heading("Transcript", 2)
    doc.add_paragraph("{{diarized_transcript}}")

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    log.info("Template written to %s", dest)
    return True


# ── PDF ────────────────────────────────────────────────────────────────────

def _html_template(config) -> str:
    custom = config and getattr(config, "output_template_pdf", "")
    if custom:
        p = Path(custom)
        if p.exists():
            return p.read_text(encoding="utf-8")
    user_default = Path.home() / ".whisperapp" / "templates" / "default.html"
    if user_default.exists():
        return user_default.read_text(encoding="utf-8")
    pkg = _PKG_TEMPLATES / "default.html"
    if pkg.exists():
        return pkg.read_text(encoding="utf-8")
    return ""


def _build_html_ctx(ctx: dict) -> dict:
    """Extend the base context with HTML-specific rendered blocks."""
    extra = dict(ctx)

    # Transcript block — render segments as rows, pause markers as italic spans
    rows = []
    for seg in ctx["segments"].split("\n"):
        seg = seg.strip()
        if not seg:
            continue
        if seg.startswith("[") and seg.endswith("]") and "→" not in seg:
            rows.append(
                f'<div class="transcript-line pause-marker">'
                f'<span></span><span>{seg}</span></div>'
            )
        elif "→" in seg:
            ts, _, rest = seg.partition("  ")
            rows.append(
                f'<div class="transcript-line">'
                f'<span class="transcript-ts">{ts.strip()}</span>'
                f'<span>{rest.strip()}</span></div>'
            )
        else:
            rows.append(
                f'<div class="transcript-line">'
                f'<span></span><span>{seg}</span></div>'
            )
    extra["transcript_block"] = "\n".join(rows) if rows else (
        "<p>" + ctx["diarized_transcript"].replace("\n", "<br>") + "</p>"
    )

    # Meeting notes block
    if ctx.get("meeting_notes"):
        notes_html = ctx["meeting_notes"].replace("\n", "<br>")
        extra["meeting_notes_block"] = (
            f'<h2>Meeting Notes</h2>'
            f'<div class="notes-box">{notes_html}</div>'
        )
    else:
        extra["meeting_notes_block"] = ""

    # Legal provenance block
    size_str = ctx.get("file_size_bytes", "—")
    try:
        size_str = f"{int(size_str):,} bytes"
    except (ValueError, TypeError):
        pass
    extra["provenance_block"] = (
        f"TRANSCRIPT OF AUDIO RECORDING\n"
        f"{'═' * 52}\n"
        f"Source file  : {ctx.get('filename', '—')}\n"
        f"Source path  : {ctx.get('filepath', '—')}\n"
        f"SHA-256      : {ctx.get('source_hash', '—')}\n"
        f"File size    : {size_str}\n"
        f"Duration     : {ctx.get('duration', '—')}\n"
        f"Transcribed  : {ctx.get('datetime', '—')}\n"
        f"Model        : {ctx.get('model', '—')}\n"
        f"Language     : {ctx.get('language', '—')}\n"
        f"{'═' * 52}"
    )

    # Legal block — HTML version of the line-numbered transcript
    legal_rows = []
    for line in ctx["legal_transcript"].split("\n"):
        if not line.strip():
            continue
        parts = line.split("  ", 3)
        if len(parts) == 4:
            num, ts, speaker, text = parts
            is_pause = ts.strip() == "" and speaker.strip() == ""
            if is_pause:
                legal_rows.append(
                    f'<div class="ll-row ll-pause">'
                    f'<span class="ll-num">{num}</span>'
                    f'<span class="ll-ts"></span>'
                    f'<span class="ll-speaker"></span>'
                    f'<span class="ll-text">{text.strip()}</span>'
                    f'</div>'
                )
            else:
                legal_rows.append(
                    f'<div class="ll-row">'
                    f'<span class="ll-num">{num}</span>'
                    f'<span class="ll-ts">{ts.strip()}</span>'
                    f'<span class="ll-speaker">{speaker.strip()}</span>'
                    f'<span class="ll-text">{text.strip()}</span>'
                    f'</div>'
                )
        else:
            legal_rows.append(f'<div class="ll-row"><span class="ll-num"></span><span class="ll-text" style="grid-column:2/-1">{line.strip()}</span></div>')
    extra["legal_block"] = "\n".join(legal_rows)

    return extra


def write_pdf(result: dict, source_file: str, out_path: Path, config=None):
    try:
        from weasyprint import HTML
    except ImportError:
        log.warning("weasyprint not installed — skipping PDF. "
                    "Install with: pip install weasyprint")
        return

    tmpl = _html_template(config)
    if not tmpl:
        log.warning("No HTML template found for PDF output")
        return

    ctx = _build_html_ctx(build_context(result, source_file))
    html_str = substitute(tmpl, ctx)
    HTML(string=html_str).write_pdf(str(out_path))
    log.info("Wrote PDF: %s", out_path)
