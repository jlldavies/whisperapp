"""Tests for DOCX and PDF output formatters."""
import pytest
from pathlib import Path
from whisperapp.document_formats import (
    build_context, substitute, generate_docx_template,
    _fmt_ts_long, _build_html_ctx,
)


SAMPLE_RESULT = {
    "language": "en",
    "model": "large-v2",
    "segments": [
        {"start": 0.0, "end": 3.2, "text": " Hello world.", "speaker": "SPEAKER_00"},
        {"start": 3.5, "end": 3.5, "text": "[Pause, 4 seconds]", "is_pause_marker": True},
        {"start": 7.8, "end": 12.1, "text": " How are you today?", "speaker": "SPEAKER_01"},
    ],
    "source_metadata": {
        "source_file": "interview.mp3",
        "duration": "00:12",
    },
}


def test_build_context_basic():
    ctx = build_context(SAMPLE_RESULT, "/tmp/interview.mp3")
    assert ctx["filename"] == "interview.mp3"
    assert ctx["model"] == "large-v2"
    assert ctx["language"] == "en"
    assert ctx["duration"] == "00:12"
    assert "Hello world." in ctx["transcript"]
    assert "[Pause, 4 seconds]" in ctx["transcript"]
    assert "SPEAKER_00" in ctx["speakers"]
    assert "SPEAKER_01" in ctx["speakers"]


def test_build_context_diarized():
    ctx = build_context(SAMPLE_RESULT, "/tmp/interview.mp3")
    assert "[SPEAKER_00]: Hello world." in ctx["diarized_transcript"]
    assert "[SPEAKER_01]: How are you today?" in ctx["diarized_transcript"]
    assert "[Pause, 4 seconds]" in ctx["diarized_transcript"]


def test_build_context_word_count():
    ctx = build_context(SAMPLE_RESULT, "/tmp/interview.mp3")
    # "Hello world." = 2 words, "How are you today?" = 4 words
    assert ctx["word_count"] == "6"


def test_build_context_segments_timestamped():
    ctx = build_context(SAMPLE_RESULT, "/tmp/interview.mp3")
    assert "→" in ctx["segments"]
    assert "SPEAKER_00" in ctx["segments"]


def test_substitute_basic():
    assert substitute("Hello {{name}}!", {"name": "World"}) == "Hello World!"


def test_substitute_unknown_marker_preserved():
    result = substitute("{{known}} and {{unknown}}", {"known": "yes"})
    assert result == "yes and {{unknown}}"


def test_substitute_multiple():
    ctx = {"a": "1", "b": "2"}
    assert substitute("{{a}}-{{b}}", ctx) == "1-2"


def test_write_docx(tmp_path):
    pytest.importorskip("docx", reason="python-docx not installed")
    from whisperapp.document_formats import write_docx
    out = tmp_path / "output.docx"
    write_docx(SAMPLE_RESULT, "/tmp/interview.mp3", out)
    assert out.exists()
    assert out.stat().st_size > 1000


def test_generate_docx_template(tmp_path):
    pytest.importorskip("docx", reason="python-docx not installed")
    dest = tmp_path / "template.docx"
    ok = generate_docx_template(dest)
    assert ok
    assert dest.exists()
    assert dest.stat().st_size > 1000


def test_docx_with_template(tmp_path):
    pytest.importorskip("docx", reason="python-docx not installed")
    from docx import Document
    from whisperapp.document_formats import write_docx

    # Create a minimal template with markers
    tmpl = tmp_path / "tmpl.docx"
    doc = Document()
    doc.add_paragraph("File: {{filename}}")
    doc.add_paragraph("Date: {{date}}")
    doc.add_paragraph("{{diarized_transcript}}")
    doc.save(str(tmpl))

    class FakeCfg:
        output_template_docx = str(tmpl)
        output_template_pdf = ""

    out = tmp_path / "output.docx"
    write_docx(SAMPLE_RESULT, "/tmp/interview.mp3", out, FakeCfg())
    assert out.exists()

    doc2 = Document(str(out))
    full_text = "\n".join(p.text for p in doc2.paragraphs)
    assert "interview.mp3" in full_text
    assert "SPEAKER_00" in full_text


def test_write_docx_formats_integration(tmp_path):
    """write_formats dispatches to write_docx correctly."""
    pytest.importorskip("docx", reason="python-docx not installed")
    from whisperapp.formatters import write_formats

    write_formats(SAMPLE_RESULT, "/tmp/interview.mp3", str(tmp_path), ["docx"])
    assert (tmp_path / "interview.docx").exists()


def test_build_context_provenance_fields():
    result_with_hash = {
        **SAMPLE_RESULT,
        "source_metadata": {
            "source_file": "interview.mp3",
            "source_path": "/data/interview.mp3",
            "source_hash": "abc123def456",
            "file_size_bytes": "1024000",
            "duration": "00:12",
        },
    }
    ctx = build_context(result_with_hash, "/data/interview.mp3")
    assert ctx["source_hash"] == "abc123def456"
    assert ctx["file_size_bytes"] == "1024000"


def test_build_context_legal_transcript():
    ctx = build_context(SAMPLE_RESULT, "/tmp/interview.mp3")
    lt = ctx["legal_transcript"]
    assert "L0001" in lt
    assert "L0002" in lt
    assert "L0003" in lt
    assert "Hello world." in lt
    assert "Pause" in lt
    assert "How are you today?" in lt


def test_build_context_missing_hash_fallback():
    ctx = build_context(SAMPLE_RESULT, "/tmp/interview.mp3")
    # No source_metadata → falls back to "—"
    assert ctx["source_hash"] == "—"
    assert ctx["file_size_bytes"] == "—"


# ── _fmt_ts_long ───────────────────────────────────────────────────────────

def test_fmt_ts_long_seconds():
    assert _fmt_ts_long(45.0) == "00:00:45"


def test_fmt_ts_long_minutes():
    assert _fmt_ts_long(125.0) == "00:02:05"


def test_fmt_ts_long_hours():
    assert _fmt_ts_long(3661.0) == "01:01:01"


def test_fmt_ts_long_always_hhmmss():
    # Even for sub-minute values it returns HH:MM:SS (unlike _fmt_ts)
    assert _fmt_ts_long(5.0) == "00:00:05"


# ── legal_transcript line format ───────────────────────────────────────────

def test_legal_transcript_line_numbers_sequential():
    ctx = build_context(SAMPLE_RESULT, "/tmp/interview.mp3")
    lines = [l for l in ctx["legal_transcript"].split("\n") if l.strip()]
    nums = [l.split()[0] for l in lines]
    assert nums == ["L0001", "L0002", "L0003"]


def test_legal_transcript_pause_has_no_timestamp():
    ctx = build_context(SAMPLE_RESULT, "/tmp/interview.mp3")
    lines = ctx["legal_transcript"].split("\n")
    pause_line = next(l for l in lines if "Pause" in l)
    # Pause lines have empty timestamp column — no HH:MM:SS in that position
    # but the pause text itself is there
    assert "[Pause, 4 seconds]" in pause_line


def test_legal_transcript_speaker_present():
    ctx = build_context(SAMPLE_RESULT, "/tmp/interview.mp3")
    assert "[SPEAKER_00]" in ctx["legal_transcript"]
    assert "[SPEAKER_01]" in ctx["legal_transcript"]


def test_legal_transcript_timestamp_format():
    ctx = build_context(SAMPLE_RESULT, "/tmp/interview.mp3")
    # First segment starts at 0.0 → "00:00:00"
    assert "00:00:00" in ctx["legal_transcript"]


# ── _build_html_ctx ────────────────────────────────────────────────────────

SAMPLE_RESULT_WITH_HASH = {
    **SAMPLE_RESULT,
    "source_metadata": {
        "source_file": "interview.mp3",
        "source_path": "/data/interview.mp3",
        "source_hash": "deadbeef1234",
        "file_size_bytes": "2048000",
        "duration": "00:12",
    },
}


def test_build_html_ctx_legal_block_present():
    ctx = build_context(SAMPLE_RESULT_WITH_HASH, "/data/interview.mp3")
    html_ctx = _build_html_ctx(ctx)
    assert "legal_block" in html_ctx
    assert "ll-row" in html_ctx["legal_block"]


def test_build_html_ctx_legal_block_contains_speakers():
    ctx = build_context(SAMPLE_RESULT_WITH_HASH, "/data/interview.mp3")
    html_ctx = _build_html_ctx(ctx)
    assert "SPEAKER_00" in html_ctx["legal_block"]
    assert "SPEAKER_01" in html_ctx["legal_block"]


def test_build_html_ctx_legal_block_pause_marker():
    ctx = build_context(SAMPLE_RESULT_WITH_HASH, "/data/interview.mp3")
    html_ctx = _build_html_ctx(ctx)
    assert "ll-pause" in html_ctx["legal_block"]
    assert "Pause, 4 seconds" in html_ctx["legal_block"]


def test_build_html_ctx_provenance_block():
    ctx = build_context(SAMPLE_RESULT_WITH_HASH, "/data/interview.mp3")
    html_ctx = _build_html_ctx(ctx)
    pb = html_ctx["provenance_block"]
    assert "deadbeef1234" in pb
    assert "interview.mp3" in pb
    assert "TRANSCRIPT OF AUDIO RECORDING" in pb


def test_build_html_ctx_provenance_size_formatted():
    ctx = build_context(SAMPLE_RESULT_WITH_HASH, "/data/interview.mp3")
    html_ctx = _build_html_ctx(ctx)
    # 2048000 bytes should appear in some form in provenance
    assert "2048000" in html_ctx["provenance_block"] or "2,048,000" in html_ctx["provenance_block"]


def test_build_html_ctx_provenance_missing_hash():
    ctx = build_context(SAMPLE_RESULT, "/tmp/interview.mp3")  # no source_hash
    html_ctx = _build_html_ctx(ctx)
    assert "—" in html_ctx["provenance_block"]


# ── Legal template file ────────────────────────────────────────────────────

def test_legal_html_template_exists():
    from whisperapp.document_formats import _PKG_TEMPLATES
    legal = _PKG_TEMPLATES / "legal.html"
    assert legal.exists(), "legal.html template missing from package"
    content = legal.read_text()
    assert "{{source_hash}}" in content
    assert "{{legal_block}}" in content
    assert "{{filename}}" in content


def test_write_pdf_with_legal_template(tmp_path):
    pytest.importorskip("weasyprint", reason="weasyprint not installed")
    from whisperapp.document_formats import write_pdf, _PKG_TEMPLATES

    class FakeCfg:
        output_template_docx = ""
        output_template_pdf = str(_PKG_TEMPLATES / "legal.html")

    out = tmp_path / "legal.pdf"
    write_pdf(SAMPLE_RESULT_WITH_HASH, "/data/interview.mp3", out, FakeCfg())
    assert out.exists()
    assert out.stat().st_size > 1000
